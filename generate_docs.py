"""
generate_docs.py
================
Generates a comprehensive DOCX documentation file for the
Options Pricing Engine project — line-by-line code explanations,
architecture overview, limitations, advantages, and improvement ideas.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_background(cell, hex_color):
    """Set background colour of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p


def add_para(doc, text, bold=False, italic=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p


def add_code_block(doc, code_text):
    """Add a styled monospace code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # light grey background via style
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1f, 0x69, 0x8a)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def section_divider(doc):
    doc.add_paragraph()
    p = doc.add_paragraph('─' * 80)
    p.runs[0].font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    p.runs[0].font.size = Pt(8)
    doc.add_paragraph()


# ─── Build Document ───────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── Default paragraph font ────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ══════════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════════════════════════════════
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run('Options Pricing Engine')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_p.add_run('Complete Project Documentation')
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    sub_p2 = doc.add_paragraph()
    sub_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = sub_p2.add_run(
        'Line-by-line Code Explanation · Architecture · '
        'Limitations · Advantages · Improvement Roadmap'
    )
    run3.font.size = Pt(12)
    run3.italic = True
    run3.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Author / Date ─────────────────────────────────────────────────────────
    auth_p = doc.add_paragraph()
    auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth_run = auth_p.add_run('Chaithanya  ·  2026')
    auth_run.font.size = Pt(12)
    auth_run.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, 'Table of Contents', level=1)
    toc_items = [
        ('1',  'Project Overview'),
        ('2',  'Repository Structure — Every File Explained'),
        ('3',  'src/pricers/bs.py  — Black-Scholes Pricer (Line by Line)'),
        ('4',  'src/greeks/greeks.py  — Analytical Greeks (Line by Line)'),
        ('5',  'src/greeks/mlp_greeks.py  — Neural Greeks (Line by Line)'),
        ('6',  'src/pricers/monte_carlo.py  — Monte Carlo Pricer (Line by Line)'),
        ('7',  'src/pricers/binomial.py  — CRR Binomial Tree (Line by Line)'),
        ('8',  'src/pricers/heston.py  — Heston Stochastic Volatility (Line by Line)'),
        ('9',  'src/pricers/neural_net.py  — MLP Pricer (Line by Line)'),
        ('10', 'src/pricers/lstm.py  — LSTM Volatility Forecaster (Line by Line)'),
        ('11', 'src/pricers/vae.py  — Variational Autoencoder (Line by Line)'),
        ('12', 'src/utils/data.py  — Data Fetching Utilities (Line by Line)'),
        ('13', 'src/utils/interpolate.py  — IV Surface Interpolation (Line by Line)'),
        ('14', 'app.py  — Streamlit Dashboard (Line by Line)'),
        ('15', 'tests/test_greeks.py  — Unit Tests (Line by Line)'),
        ('16', 'scripts/ — Helper Scripts Explained'),
        ('17', 'How Everything Connects — End-to-End Data Flow'),
        ('18', 'Mathematical Foundations'),
        ('19', 'Empirical Results & Benchmarks'),
        ('20', 'Advantages of the Project'),
        ('21', 'Limitations of the Project'),
        ('22', 'How to Improve the Project — Roadmap'),
    ]
    for num, item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        r1 = p.add_run(f'{num}.  ')
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(item)
        r2.font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — PROJECT OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '1. Project Overview', level=1)
    add_para(doc,
        'The Options Pricing Engine is a research-grade Python project that '
        'implements, compares, and benchmarks seven distinct methodologies for '
        'pricing European and American options. The goal is to systematically '
        'evaluate classical financial models alongside modern machine-learning '
        'approaches on 200 real SPY (S&P 500 ETF) options.',
    )
    doc.add_paragraph()
    add_para(doc, 'The seven pricing methods are:', bold=True)
    methods = [
        ('Black-Scholes (BS)', 'Closed-form analytical solution for European options under constant volatility.'),
        ('CRR Binomial Tree', 'Lattice model — prices European and American options; handles early exercise.'),
        ('Monte Carlo (Basic, Antithetic, Control Variate)', 'Simulation-based approach — flexible for exotic and path-dependent options.'),
        ('Heston Stochastic Volatility', 'Monte Carlo under a two-factor model that reproduces the real volatility smile.'),
        ('MLP Neural Network Pricer', 'Multi-layer perceptron trained on 100,000 synthetic BS prices — ultra-fast inference.'),
        ('LSTM Volatility Forecaster', 'Sequence model that predicts next-day realized volatility, fed into BS.'),
        ('VAE Implied Volatility Surface', 'Variational autoencoder learns the latent structure of the IV surface for interpolation.'),
    ]
    for name, desc in methods:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        r = p.add_run(f'• {name}: ')
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(desc).font.size = Pt(11)

    doc.add_paragraph()
    add_para(doc,
        'The project is structured as a Python package (src/), a Streamlit '
        'interactive dashboard (app.py), unit tests (tests/), and helper '
        'scripts (scripts/). It also includes a full academic paper '
        '(options-pricing-ml-Chaithanya-2026.pdf).',
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — REPOSITORY STRUCTURE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '2. Repository Structure — Every File Explained', level=1)
    add_para(doc,
        'Below is every file in the project, why it exists, and what it does.',
        bold=False
    )
    doc.add_paragraph()

    files = [
        # ── Root ──────────────────────────────────────────────────────────────
        ('app.py', 'ROOT',
         '1,739-line Streamlit web application. Provides the interactive '
         'pricing dashboard with 4 tabs: Pricing, Greeks, Vol Smile, and '
         'Model Comparison. Imports all pricers and renders live Plotly charts.'),
        ('requirements.txt', 'ROOT',
         'Pinned list of all Python package dependencies (numpy, torch, '
         'streamlit, plotly, yfinance, scipy, scikit-learn, etc.) needed '
         'to reproduce the environment exactly.'),
        ('README.md', 'ROOT',
         'Project landing page. Contains mathematical equations for BS, MC, '
         'and Heston; the full empirical results table (7 models × 7 metrics); '
         'quick-start installation instructions.'),
        ('fix_imports.py', 'ROOT',
         'One-off utility script that patches import statements across the '
         'codebase. Not part of runtime — only run once during refactoring.'),
        ('insert_images.py', 'ROOT',
         'Helper script that programmatically inserts generated PNG charts '
         'into the Word or PDF paper document.'),
        ('paper_academic.md', 'ROOT',
         'Academic paper draft in Markdown format. Describes methodology, '
         'results, and conclusions in IEEE/journal style.'),
        ('paper_humanized.md', 'ROOT',
         'A more readable, plain-language version of the paper draft.'),
        ('paper_draft.md', 'ROOT',
         'Earlier working draft of the paper.'),
        ('options-pricing-ml-Chaithanya-2026.pdf', 'ROOT',
         'Final compiled academic paper PDF submitted to SSRN.'),
        ('options-pricing-ml-Chaithanya-2026.tex', 'ROOT',
         'LaTeX source file for the academic paper.'),
        ('results_detailed.csv', 'ROOT',
         '200-option SPY test dataset. Each row is one real option with '
         'Strike, Market price, IV, moneyness bucket, spot price, and T.'),
        ('test_options.csv', 'ROOT',
         'Smaller test CSV used during development and scripted evaluations.'),
        ('comparison_table.csv', 'ROOT',
         'Summary table (7 rows × 7 cols) of model benchmarks: MAE, RMSE, '
         '% within 5%, ATM MAE, OTM MAE, Speed. Used in the paper.'),
        ('mlp_pricer.pth', 'ROOT',
         'Saved PyTorch state dict of the trained MLP model (184 KB).'),
        ('vae_weights.pth', 'ROOT',
         'Saved PyTorch state dict of the trained VAE model (51 KB).'),
        ('iv_surfaces.pt', 'ROOT',
         'PyTorch tensor of 500 implied-volatility surfaces (5 expiries × '
         '10 strikes each) used to train the VAE.'),
        # ── PNGs ──────────────────────────────────────────────────────────────
        ('greeks_dashboard.png', 'ROOT (output)',
         '5-panel Greeks visualization — Delta, Gamma, Vega, Theta, Rho vs '
         'various parameters. Generated by bs.py __main__ block.'),
        ('mc_analysis.png', 'ROOT (output)',
         '4-panel Monte Carlo variance reduction analysis — convergence, '
         'std error, ratio, sampling distribution.'),
        ('crr_benchmark.png', 'ROOT (output)',
         'CRR binomial tree convergence to Black-Scholes across N steps.'),
        ('heston_smile.png', 'ROOT (output)',
         'Heston implied-volatility smile (curved) vs flat Black-Scholes '
         'smile — the key result demonstrating stochastic vol advantage.'),
        ('neural_vs_bs_delta.png', 'ROOT (output)',
         'Neural Delta vs BS Delta across moneyness range.'),
        ('mlp_loss_curve.png', 'ROOT (output)',
         'MLP training and validation loss over 100 epochs.'),
        ('loss_curve.png', 'ROOT (output)',
         'LSTM training/validation loss over 50 epochs.'),
        ('historical_volatility.png', 'ROOT (output)',
         '30-day realized volatility of SPY plotted over time.'),
        ('volatility_forecast_comparison.png', 'ROOT (output)',
         'LSTM predicted volatility vs actual realized volatility.'),
        ('interpolation_comparison.png', 'ROOT (output)',
         '3-panel: original IV surface | linear interpolation | VAE interpolation.'),
        ('pricer_comparison.png', 'ROOT (output)',
         'Bar-chart comparison of all model prices for a standard option.'),
        ('generated_surfaces.png', 'ROOT (output)',
         'Multiple VAE-generated IV surfaces showing learned diversity.'),
        # ── src/pricers ───────────────────────────────────────────────────────
        ('src/pricers/bs.py', 'src/pricers',
         'Black-Scholes closed-form pricer. Implements norm_cdf (hand-coded '
         'Horner rational approximation), norm_pdf, black_scholes(), '
         'and implied_vol() (Newton-Raphson with Brent fallback).'),
        ('src/pricers/binomial.py', 'src/pricers',
         'CRR binomial lattice. Implements crr_parameters(), '
         'build_stock_tree() (vectorised NumPy), price_option() with backward '
         'induction for European and American exercises, numerical Greeks, '
         'and a convergence plot generator.'),
        ('src/pricers/monte_carlo.py', 'src/pricers',
         'Monte Carlo pricer with three variance-reduction techniques: basic '
         'mc_price(), mc_antithetic(), mc_control_variate(). Also includes '
         'mc_barrier_option() for exotic options and a 4-panel analysis plot.'),
        ('src/pricers/heston.py', 'src/pricers',
         'Heston (1993) stochastic volatility model priced via Euler-Maruyama '
         'Monte Carlo with antithetic variates. Includes smile computation, '
         'calibration via Nelder-Mead, and SPY test-set evaluation.'),
        ('src/pricers/neural_net.py', 'src/pricers',
         'MLP-based option pricer. Generates 100,000 synthetic training '
         'samples via BS, trains a 4-layer MLP (256→128→64→1), evaluates '
         'on in-distribution and OOD test cases, and tests on real SPY data.'),
        ('src/pricers/lstm.py', 'src/pricers',
         'LSTM volatility forecaster. Downloads 5 years of SPY data, '
         'computes rolling realized volatility, trains an LSTM to predict '
         'next-day vol, then feeds LSTM output into Black-Scholes.'),
        ('src/pricers/vae.py', 'src/pricers',
         'Variational Autoencoder for IV surface generation. Encoder '
         '(50→64→32→μ/σ), decoder (8→32→64→50), reparameterization trick, '
         'and ELBO loss = MSE reconstruction + β·KL divergence.'),
        # ── src/greeks ────────────────────────────────────────────────────────
        ('src/greeks/greeks.py', 'src/greeks',
         'Analytical Black-Scholes Greeks: Delta, Gamma, Vega, Theta, Rho — '
         'all derived from the standard d1/d2 formulas. Imports norm_cdf and '
         'norm_pdf from bs.py.'),
        ('src/greeks/mlp_greeks.py', 'src/greeks',
         'Computes neural network Greeks (Delta, Vega) via torch.autograd '
         'automatic differentiation. Compares neural Greeks to analytical BS '
         'Greeks and generates the neural_vs_bs_delta.png plot.'),
        # ── src/utils ─────────────────────────────────────────────────────────
        ('src/utils/data.py', 'src/utils',
         'Data pipeline. Fetches implied-volatility surfaces for up to 300 '
         'S&P 500 tickers via yfinance, interpolates sparse data to a '
         '5×10 grid, pads with synthetic SABR-like surfaces to 500 total, '
         'and saves as iv_surfaces.pt.'),
        ('src/utils/interpolate.py', 'src/utils',
         'Compares two IV surface imputation strategies: (1) scipy griddata '
         'linear interpolation and (2) VAE latent-space optimization. '
         'Generates a 3-panel comparison plot.'),
        # ── tests ─────────────────────────────────────────────────────────────
        ('tests/test_greeks.py', 'tests',
         'unittest TestCase with 6 tests: delta bounds, gamma ATM-peak, '
         'vega positivity, theta negativity (time decay), rho sign '
         'conventions, and implied_vol() return type check.'),
        # ── scripts ───────────────────────────────────────────────────────────
        ('scripts/train.py', 'scripts',
         'Short entry-point that calls generate_synthetic_data() and '
         'train_model() from neural_net.py to train and save mlp_pricer.pth.'),
        ('scripts/generate.py', 'scripts',
         'Calls data.py main() to fetch/synthesize IV surfaces and saves '
         'iv_surfaces.pt.'),
        ('scripts/benchmark_crr.py', 'scripts',
         'Runs CRR binomial tree at N=10 to 500 and plots convergence to '
         'Black-Scholes, saving crr_benchmark.png.'),
        ('scripts/pricer_comparison.py', 'scripts',
         'Runs all six pricers on a standard ATM option (S=K=100, T=1, '
         'r=5%, σ=20%) and generates a comparison bar chart.'),
        ('scripts/volatility_smile.py', 'scripts',
         'Entry-point for the Heston smile plot: calls plot_heston_smile() '
         'and saves heston_smile.png.'),
        ('scripts/volatility_surface.py', 'scripts',
         'Generates and visualizes 3D volatility surfaces (multiple expiries).'),
        ('scripts/eval_real.py', 'scripts',
         'Evaluates all pricers on real SPY options from results_detailed.csv '
         'and computes MAE, RMSE, and % within 5% for each model.'),
        ('scripts/dm_test.py', 'scripts',
         'Diebold-Mariano statistical test — checks whether one model\'s '
         'forecast errors are significantly different from another\'s.'),
        ('scripts/run_definitive_experiment.py', 'scripts',
         'Master orchestration script that runs the full evaluation pipeline: '
         'all seven models × 200 SPY options, writes results_detailed.csv, '
         'comparison_table.csv, and all paper figures.'),
        ('scripts/generate_pdf.py', 'scripts',
         'Programmatically assembles the academic paper PDF from markdown '
         'content and inserts generated PNG figures.'),
        ('scripts/generate_html_pdf.py', 'scripts',
         'Alternative paper generator that converts markdown to HTML then to '
         'PDF via weasyprint or equivalent.'),
    ]

    for fname, location, description in files:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.0)
        r1 = p.add_run(f'📄 {fname}')
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.4)
        p2.add_run(description).font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — bs.py LINE BY LINE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '3. src/pricers/bs.py — Black-Scholes Pricer (Line by Line)', level=1)
    add_para(doc,
        'This is the foundational file. Every other pricer either calls '
        'black_scholes() for comparison or imports norm_cdf / norm_pdf.',
        italic=True
    )
    doc.add_paragraph()

    bs_lines = [
        ('1',  'import math',
         'Imports Python\'s standard math module for log, exp, sqrt, pi — '
         'used throughout every formula.'),
        ('2',  'import matplotlib.pyplot as plt',
         'Matplotlib is imported for the Greeks dashboard plot generation '
         'at the bottom of the file (in __main__).'),
        ('3',  'import scipy.optimize',
         'SciPy\'s optimize module provides brentq() — Brent\'s root-finding '
         'method — used as fallback in implied_vol() when Newton-Raphson fails.'),
        ('5',  'from typing import Tuple, Optional',
         'Type hints: Tuple for functions returning (call, put), Optional for '
         'implied_vol which can return None on convergence failure.'),
        ('8–35', 'def norm_cdf(x)',
         'Cumulative Distribution Function of the standard normal — '
         'implemented from scratch using the Abramowitz & Stegun (1964) '
         'polynomial approximation 26.2.17, accurate to ~1×10⁻⁷.\n'
         '  Line 20: x_abs = abs(x)  — work with positive x for symmetry.\n'
         '  Line 22: t = 1/(1 + 0.2316419·|x|)  — transform for Horner scheme.\n'
         '  Lines 25–29: polynomial coefficients a1…a5.\n'
         '  Line 31: poly = Horner evaluation of the 5th-degree polynomial.\n'
         '  Line 32: pdf = standard normal PDF value at |x|.\n'
         '  Line 34: cdf_positive = 1 - pdf·poly  (for x ≥ 0).\n'
         '  Line 35: returns cdf_positive or 1-cdf_positive using symmetry N(-x) = 1-N(x).'),
        ('38–47', 'def norm_pdf(x)',
         'Probability Density Function: (1/√(2π)) · exp(-x²/2).\n'
         'Used in: vega = S·N\'(d1)·√T, gamma = N\'(d1)/(S·σ·√T), theta.'),
        ('50–75', 'def black_scholes(S, K, T, r, sigma)',
         'Core Black-Scholes pricing function.\n'
         '  Line 66: d1 = [ln(S/K) + (r + σ²/2)·T] / (σ·√T)\n'
         '           — risk-adjusted log-moneyness, scaled by σ√T.\n'
         '  Line 67: d2 = d1 - σ·√T  — probability-adjusted d1.\n'
         '  Line 70: pv_K = K·e^(-rT)  — present value of strike.\n'
         '  Line 72: call = S·N(d1) - K·e^(-rT)·N(d2)\n'
         '           — risk-neutral expected payoff, discounted.\n'
         '  Line 73: put = K·e^(-rT)·N(-d2) - S·N(-d1)\n'
         '           — derived from put-call parity.'),
        ('78–138', 'def implied_vol(market_price, S, K, T, r, option_type)',
         'Inverts the BS formula to find the σ that makes BS price match market.\n'
         '  Line 104: sigma = 0.2  — initial guess (20% vol).\n'
         '  Lines 106–126: Newton-Raphson iteration:\n'
         '    σ_{n+1} = σ_n - (BS(σ_n) - market_price) / vega(σ_n)\n'
         '    Stops when |price_diff| < 0.0001 or after 100 iterations.\n'
         '  Lines 123–124: Clamps σ > 0 to avoid negative volatility.\n'
         '  Lines 129–138: Fallback to Brent\'s method on [0.001, 10] if\n'
         '    Newton-Raphson fails to converge (e.g., deep ITM/OTM options).'),
        ('142–253', '__main__ block',
         'Runs when the file is executed directly (python -m src.pricers.bs).\n'
         '  Lines 143–155: Verify pricing and put-call parity for ATM option.\n'
         '  Lines 165–177: Verify all 5 Greeks at S=100.\n'
         '  Lines 180–253: Generates the 5-panel greeks_dashboard.png:\n'
         '    Subplot 1: Delta vs Spot (call ↑ from 0 to 1, put ↓ from -1 to 0)\n'
         '    Subplot 2: Gamma vs Spot (peaks at ATM, symmetric bell)\n'
         '    Subplot 3: Vega vs Volatility (increases with vol)\n'
         '    Subplot 4: Theta vs Time (negative, accelerates near expiry)\n'
         '    Subplot 5: Rho vs Interest Rate (positive for calls, negative for puts)'),
    ]

    for lines, code, explanation in bs_lines:
        p = doc.add_paragraph()
        r1 = p.add_run(f'Lines {lines}: ')
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        r1.font.size = Pt(11)
        r2 = p.add_run(code)
        r2.font.name = 'Courier New'
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)

        for line in explanation.split('\n'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.4)
            p2.add_run(line).font.size = Pt(11)

        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 4 — greeks.py LINE BY LINE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '4. src/greeks/greeks.py — Analytical Greeks (Line by Line)', level=1)
    add_para(doc,
        'Five pure-math functions, each returning the sensitivity of the option '
        'price to one input variable.',
        italic=True
    )
    doc.add_paragraph()

    greeks_lines = [
        ('1',  'import math', 'Standard library math functions.'),
        ('2',  'from typing import Tuple', 'Type hint for functions that return (call_greek, put_greek).'),
        ('3',  'from src.pricers.bs import norm_cdf, norm_pdf',
         'Reuses the hand-coded CDF and PDF from bs.py — avoids duplication.'),
        ('6–27', 'def black_scholes_delta(S, K, T, r, sigma)',
         'Delta = ∂V/∂S — how much option price changes per $1 move in stock.\n'
         '  Line 24: d1 = same formula as in bs.py.\n'
         '  Line 25: call_delta = N(d1)  — ranges from 0 (deep OTM) to 1 (deep ITM).\n'
         '  Line 26: put_delta  = N(d1) - 1  — same as call_delta - 1 (put-call parity).'),
        ('30–48', 'def black_scholes_gamma(S, K, T, r, sigma)',
         'Gamma = ∂²V/∂S² = ∂Delta/∂S — curvature of price w.r.t. stock.\n'
         '  Line 47: gamma = N\'(d1) / (S · σ · √T)\n'
         '  Identical for calls and puts. Peaks when S = K (ATM). Positive always.'),
        ('51–69', 'def black_scholes_vega(S, K, T, r, sigma)',
         'Vega = ∂V/∂σ — sensitivity to a 1-unit change in volatility.\n'
         '  Line 68: vega = S · N\'(d1) · √T\n'
         '  Always positive. Largest ATM. Decreases far ITM or OTM.'),
        ('72–99', 'def black_scholes_theta(S, K, T, r, sigma)',
         'Theta = ∂V/∂t — time decay (how much the option loses per day).\n'
         '  Line 90–91: d1, d2 as usual.\n'
         '  Line 93: first_term = -(S · N\'(d1) · σ) / (2√T)  — diffusion decay.\n'
         '  Line 94: second_term_call = r · K · e^(-rT) · N(d2)  — interest rate component.\n'
         '  Lines 97–98: call and put thetas differ by sign of second_term.\n'
         '  Always negative for long calls and puts in normal markets.'),
        ('102–125', 'def black_scholes_rho(S, K, T, r, sigma)',
         'Rho = ∂V/∂r — sensitivity to the risk-free interest rate.\n'
         '  Line 123: call_rho = K · T · e^(-rT) · N(d2)  — positive: higher r → calls worth more.\n'
         '  Line 124: put_rho = -K · T · e^(-rT) · N(-d2)  — negative: higher r → puts worth less.'),
    ]

    for lines, code, explanation in greeks_lines:
        p = doc.add_paragraph()
        r1 = p.add_run(f'Lines {lines}: ')
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        r1.font.size = Pt(11)
        r2 = p.add_run(code)
        r2.font.name = 'Courier New'
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
        for line in explanation.split('\n'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.4)
            p2.add_run(line).font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 5 — mlp_greeks.py
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '5. src/greeks/mlp_greeks.py — Neural Greeks (Line by Line)', level=1)
    add_para(doc,
        'Instead of deriving Greeks analytically, this module uses PyTorch\'s '
        'automatic differentiation (autograd) to compute them from the neural '
        'network pricer\'s output.',
        italic=True
    )
    doc.add_paragraph()

    mlp_greeks_lines = [
        ('1–5', 'Imports',
         'numpy, torch, matplotlib — standard.\n'
         'MLPPricer from neural_net — loads the trained model.\n'
         'black_scholes_delta, black_scholes_vega from greeks — for comparison.'),
        ('8–56', 'def compute_neural_greeks(model, S, K, T, r, sigma, device)',
         'Key idea: wrap inputs as torch tensors with requires_grad=True, then\n'
         'call backward() to get ∂price/∂S (Delta) and ∂price/∂σ (Vega).\n\n'
         '  Line 23: model.eval() — turns off BatchNorm / Dropout for inference.\n'
         '  Lines 26–31: S and sigma tensors have requires_grad=True so autograd\n'
         '    tracks operations on them. K, T, r do NOT — we don\'t need those derivatives.\n'
         '  Lines 34–38: Build the 4-element input: [M, T, r, σ] where M = ln(K/S).\n'
         '    This is the normalized input the MLP expects.\n'
         '  Lines 40–42: price = S × MLP(M, T, r, σ).\n'
         '    The MLP predicts C/S (normalized price). Multiply by S to get actual price.\n'
         '    This multiplication is tracked by autograd — so ∂price/∂S includes\n'
         '    both the "times S" factor and the implicit S-dependence through M = ln(K/S).\n'
         '  Lines 46–51: torch.autograd.grad() computes gradients of price w.r.t.\n'
         '    (S_tensor, sigma_tensor) simultaneously in one pass.\n'
         '  Lines 53–54: .item() extracts Python floats from tensors.'),
        ('59–93', 'def compare_greeks(model, device)',
         'Evaluates neural vs analytical Greeks at 3 moneyness points:\n'
         '  ATM (S=K=100): both should be ~0.63 delta, ~39 vega.\n'
         '  ITM (S=120, K=100): delta close to 1, vega lower.\n'
         '  OTM (S=80, K=100): delta close to 0, vega lower.\n'
         '  Prints a formatted table showing agreement or divergence.'),
        ('96–144', 'def plot_delta_vs_moneyness(model, device)',
         'Sweeps S from 50 to 150 (100 points), computes neural Delta and\n'
         'BS Delta at each point, plots both on the same axes.\n'
         'Saves neural_vs_bs_delta.png — used in the paper figure.'),
        ('147–164', '__main__ block',
         'Auto-detects GPU (CUDA/MPS) or falls back to CPU.\n'
         'Loads mlp_pricer.pth, runs compare_greeks() and plot_delta_vs_moneyness().'),
    ]

    for lines, code, explanation in mlp_greeks_lines:
        p = doc.add_paragraph()
        r1 = p.add_run(f'Lines {lines}: ')
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        r1.font.size = Pt(11)
        r2 = p.add_run(code)
        r2.font.name = 'Courier New'
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
        for line in explanation.split('\n'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.4)
            p2.add_run(line).font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 6 — monte_carlo.py
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '6. src/pricers/monte_carlo.py — Monte Carlo Pricer (Line by Line)', level=1)
    add_para(doc,
        '696 lines implementing three MC variants plus barrier options and a '
        'full analysis plot.',
        italic=True
    )
    doc.add_paragraph()

    mc_lines = [
        ('1–15', 'Module docstring + imports',
         'Explains the four features: basic MC, antithetic variates, control\n'
         'variates, comparison table. References Glasserman (2004) and Hull (2018).'),
        ('30–79', 'def mc_price(S0, K, T, r, sigma, n_paths, option_type, seed)',
         'Basic GBM Monte Carlo pricer:\n'
         '  Line 58: np.random.default_rng(seed) — reproducible random number generator.\n'
         '  Line 61: Z = rng.standard_normal(n_paths) — vectorized N(0,1) draws.\n'
         '  Line 62: S_T = S0 · exp((r - σ²/2)·T + σ·√T·Z)\n'
         '    This is the closed-form GBM terminal price under risk-neutral measure.\n'
         '    The factor (r - σ²/2) is the Itô correction to the drift.\n'
         '  Lines 64–67: Payoffs = max(S_T - K, 0) for call, max(K - S_T, 0) for put.\n'
         '  Line 70: price = e^(-rT) × mean(payoffs) — risk-neutral pricing equation.\n'
         '  Line 71: se = e^(-rT) × std(payoffs, ddof=1) / √N — standard error.'),
        ('85–148', 'def mc_antithetic(S0, K, T, r, sigma, n_paths, option_type, seed)',
         'Antithetic variates — halves variance at no extra cost:\n'
         '  Line 119: half = n_paths // 2 — draw only half the normals.\n'
         '  Line 120: Z ~ N(0,1) — half set of draws.\n'
         '  Line 121: Z_a = -Z — the antithetic pair (negative of each draw).\n'
         '  Lines 126–127: Simulate two stock prices per draw: S_T using Z, S_T_a using -Z.\n'
         '  Lines 136–140: paired = 0.5·(pay + pay_a) — average each antithetic pair.\n'
         '  WHY IT WORKS: Cov(f(Z), f(-Z)) < 0 for convex payoffs, so averaging\n'
         '    pair payoffs cancels much of the variance. Theoretical reduction: ~√2.'),
        ('154–218', 'def mc_control_variate(S0, K, T, r, sigma, n_paths, option_type, seed)',
         'Control variate using S_T (stock price) as the control:\n'
         '  Line 199: E_S_T = S0·e^(rT) — analytical expectation of S_T under Q.\n'
         '  Lines 202–203: beta = Cov(payoff, S_T) / Var(S_T) — optimal OLS coefficient.\n'
         '  Line 206: payoffs_cv = payoffs - β·(S_T - E[S_T])\n'
         '    Subtracts the mean-zero control to remove correlated noise component.\n'
         '  WHY IT WORKS: β captures how much of the MC error is explained by\n'
         '    S_T deviation from its mean. Removing that correlation reduces variance.'),
        ('224–349', 'def mc_barrier_option(...)',
         'Barrier option pricer — requires full path simulation:\n'
         '  Line 262: dt = T/n_steps — daily time steps (n_steps=252 by default).\n'
         '  Lines 311–316: active array tracks which paths are still "alive".\n'
         '    For "out" barriers: start True (all alive), set False on breach.\n'
         '    For "in" barriers: start False (none active), set True on breach.\n'
         '  Lines 318–329: Path loop — each step applies GBM and checks barrier.\n'
         '  Lines 336–337: payoffs = payoffs × active — knock out inactive paths.\n'
         '  WHY BARRIER NEEDS PATH SIMULATION: The barrier condition is path-dependent\n'
         '    — we must know if the stock ever touched H, not just the terminal price.'),
        ('355–401', 'def build_comparison_table(...)',
         'Runs all 3 MC methods at 3 path counts (1k, 10k, 100k).\n'
         'Computes BS reference and measures abs error, std error, and speed.'),
        ('407–616', 'def plot_mc_analysis(...)',
         '4-panel dark-theme analysis figure:\n'
         '  Panel 1: Price convergence vs N (log scale) — all 3 methods.\n'
         '  Panel 2: Standard error vs N (log-log) — shows O(1/√N) decay.\n'
         '  Panel 3: SE ratio Basic/Antithetic and Basic/CV (should be > 1).\n'
         '  Panel 4: Histogram of 500 independent MC estimates at N=10k —\n'
         '    shows antithetic has tighter distribution around the true BS price.'),
    ]

    for lines, code, explanation in mc_lines:
        p = doc.add_paragraph()
        r1 = p.add_run(f'Lines {lines}: ')
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        r1.font.size = Pt(11)
        r2 = p.add_run(code)
        r2.font.name = 'Courier New'
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
        for line in explanation.split('\n'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.4)
            p2.add_run(line).font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 7 — binomial.py
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '7. src/pricers/binomial.py — CRR Binomial Tree (Line by Line)', level=1)
    add_para(doc,
        'Cox-Ross-Rubinstein (1979) lattice model. The only pricer in the project '
        'that can handle American early-exercise options.',
        italic=True
    )
    doc.add_paragraph()

    bin_lines = [
        ('30–59', 'def crr_parameters(sigma, T, N, r)',
         'Computes the three fundamental lattice parameters:\n'
         '  dt = T/N  — length of each time step.\n'
         '  u = exp(σ·√dt)  — up factor (stock moves up by factor u each step).\n'
         '  d = 1/u  — down factor (recombining lattice: u×d = 1, so nodes recombine).\n'
         '  p = (e^(r·dt) - d) / (u - d)  — risk-neutral probability of up move.\n'
         '  Lines 53–57: Validates 0 < p < 1 — arbitrage-free condition.\n'
         '    If p ≤ 0 or p ≥ 1, the lattice is mis-parameterized.'),
        ('65–95', 'def build_stock_tree(S0, u, d, N)',
         'Builds the full (N+1)×(N+1) stock price lattice in one vectorized NumPy call:\n'
         '  Line 83–84: i_idx (time steps 0..N), j_idx (up-moves 0..N).\n'
         '  Line 88: np.meshgrid creates two (N+1)×(N+1) grids of i and j indices.\n'
         '  Line 90: S[i,j] = S0 · u^j · d^(i-j)  — node has j up-moves, (i-j) down-moves.\n'
         '  Line 93: S[J > I] = 0  — mask out physically impossible nodes (j > i).'),
        ('101–222', 'def price_option(S0, K, T, r, sigma, N, option_type, exercise)',
         'Full pricing function — forward + backward pass:\n'
         '  Line 131: crr_parameters computes u, d, p, dt.\n'
         '  Line 132: disc = e^(-r·dt)  — one-step discount factor.\n'
         '  Lines 138–143: Terminal payoffs at step N.\n'
         '  Lines 146–162: Backward induction loop from i=N-1 down to i=0:\n'
         '    V_cont[j] = disc · (p·V[j+1] + (1-p)·V[j])  — discounted expected value.\n'
         '    For American: V[j] = max(V_cont[j], intrinsic[j])  — the early-exercise check.\n'
         '    For European: V[j] = V_cont[j]  — no early exercise.\n'
         '  Line 164: price = V[0]  — root of the tree = option price at t=0.\n'
         '  Lines 168–206: Numerical Greeks:\n'
         '    Delta = (V_up - V_down) / (S_up - S_down)  — finite difference.\n'
         '    Gamma = second derivative via 3-point finite difference.\n'
         '    Theta = (price(T-1day) - price(T)) / (1/365)  — reprice for one fewer day.'),
        ('228–400', 'def plot_convergence(...)',
         'Runs price_option() for N = 10 to 500 in steps of 5.\n'
         'Generates 4-panel convergence plot:\n'
         '  Panel 1: European Call convergence to BS.\n'
         '  Panel 2: European Put convergence to BS.\n'
         '  Panel 3: American Call vs European Call — early exercise premium.\n'
         '  Panel 4: American Put vs European Put — early exercise premium.\n'
         '  All panels show the BS reference as a dashed gold line.'),
    ]

    for lines, code, explanation in bin_lines:
        p = doc.add_paragraph()
        r1 = p.add_run(f'Lines {lines}: ')
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        r1.font.size = Pt(11)
        r2 = p.add_run(code)
        r2.font.name = 'Courier New'
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
        for line in explanation.split('\n'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.4)
            p2.add_run(line).font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 8 — heston.py
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '8. src/pricers/heston.py — Heston Stochastic Volatility (Line by Line)', level=1)
    add_para(doc,
        'The most complex pricer. Heston (1993) extends Black-Scholes by letting '
        'volatility itself follow a mean-reverting stochastic process.',
        italic=True
    )
    doc.add_paragraph()

    heston_lines = [
        ('1–25', 'Module docstring',
         'Documents the two SDEs:\n'
         '  dS = r·S·dt + √v·S·dW₁   (stock price)\n'
         '  dv = κ(θ−v)dt + σᵥ·√v·dW₂  (variance, mean-reverts to θ)\n'
         '  corr(dW₁, dW₂) = ρ  — negative ρ creates the downside skew.\n'
         'Default parameters: κ=2.0, θ=0.04, σᵥ=0.3, ρ=-0.7, v₀=0.04.'),
        ('43–142', 'def heston_mc_price(...)',
         'Euler-Maruyama discretization with antithetic variates:\n'
         '  Line 91: dt = T/n_steps.\n'
         '  Line 92: sqrt_rho2 = √(1-ρ²)  — for Cholesky decomposition.\n'
         '  Lines 95–98: Initialize S, v, Sa, va — two antithetic paths each.\n'
         '  Loop (lines 100–120): For each time step:\n'
         '    Lines 102–104: Z1 ~ N(0,1), Z3 ~ N(0,1), Z2 = ρ·Z1 + √(1-ρ²)·Z3\n'
         '      This creates correlated Brownians with the correct ρ.\n'
         '    Lines 107–110: v_pos = max(v, 0) — full truncation to prevent √negative.\n'
         '    Lines 113–116: Variance update (Euler-Maruyama):\n'
         '      v_new = v + κ(θ-v)dt + σᵥ·√(v·dt)·Z2\n'
         '      v = max(v_new, 0)  — full truncation (Higham & Mao approach).\n'
         '    Lines 119–120: Log-Euler stock update:\n'
         '      S = S · exp((r - v/2)·dt + √(v·dt)·Z1)  — log-Euler guarantees S > 0.\n'
         '      Antithetic paths use -Z1 and -Z2.\n'
         '  Lines 122–131: Compute payoffs, average antithetic pairs, discount.'),
        ('148–223', 'def compute_heston_smile(...)',
         'Generates the IV smile by pricing calls at 10 strikes (80%–120% moneyness):\n'
         '  Lines 182–183: moneyness = linspace(0.80, 1.20, 10), strikes = S0 × moneyness.\n'
         '  For each K: price Heston call, then back out the BS implied vol via implied_vol().\n'
         '  Lines 208–213: Skips if Heston price ≤ intrinsic (IV extraction fails there).\n'
         '  Returns: strikes, heston_ivs, bs_flat_ivs (constant = √θ = 20%).'),
        ('229–449', 'def plot_heston_smile(...)',
         'Two-panel figure:\n'
         '  Panel 1 (IV Smile): Heston IV curved downward (left skew from ρ < 0)\n'
         '    vs flat BS smile (constant 20%). Shaded region shows the "smile premium".\n'
         '  Panel 2 (Price Curve): Heston prices vs BS prices across strikes.\n'
         '    Fills the price difference region to visualize the model divergence.'),
        ('454–468', 'def calibrate_heston(S0, K, T, r, market_price, seed)',
         'Calibrates Heston parameters to match one market option price:\n'
         '  Uses scipy Nelder-Mead optimizer (simplex method).\n'
         '  Objective = (Heston_price - market_price)².\n'
         '  n_paths=2000, n_steps=10 — deliberately fast but noisy for calibration speed.\n'
         '  Returns 5 calibrated params: [κ, θ, σᵥ, ρ, v₀].'),
        ('470–654', 'def evaluate_heston_on_test_set(csv_path, ...)',
         'Evaluates Heston on the 200 SPY options:\n'
         '  Reads results_detailed.csv, resolves column names flexibly.\n'
         '  For each option: calibrates Heston, then re-prices with n_paths=10k.\n'
         '  Buckets errors by moneyness (Deep ITM / ITM / ATM / OTM / Deep OTM).\n'
         '  Reports MAE, RMSE, % within 5%, speed per option.'),
    ]

    for lines, code, explanation in heston_lines:
        p = doc.add_paragraph()
        r1 = p.add_run(f'Lines {lines}: ')
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        r1.font.size = Pt(11)
        r2 = p.add_run(code)
        r2.font.name = 'Courier New'
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
        for line in explanation.split('\n'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.4)
            p2.add_run(line).font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 9 — neural_net.py
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '9. src/pricers/neural_net.py — MLP Pricer (Line by Line)', level=1)
    add_para(doc,
        'Trains a deep feedforward network to approximate the Black-Scholes '
        'formula — enabling microsecond inference at the cost of accuracy.',
        italic=True
    )
    doc.add_paragraph()

    nn_lines = [
        ('14–48', 'def generate_synthetic_data(num_samples=100000)',
         'Creates 100,000 random option scenarios:\n'
         '  S ∈ [50,150], K ∈ [50,150], T ∈ [0.1,2], r ∈ [0.01,0.10], σ ∈ [0.05,0.80]\n'
         '  Calls black_scholes() for each sample to get the "true" label.\n'
         '  Line 44: M = ln(K/S)  — log-moneyness: key feature for scale-invariance.\n'
         '  Line 45: norm_price = call_price/S  — normalized output.\n'
         '    WHY: The MLP learns C/S = f(M, T, r, σ), then multiply by S at runtime.\n'
         '    This makes the model scale-invariant in stock price.'),
        ('52–71', 'class MLPPricer(nn.Module)',
         '4-layer feedforward network: 4 → 256 → 128 → 64 → 1\n'
         '  Inputs: [M, T, r, σ]  (4 features)\n'
         '  Output: C/S (normalized call price)\n'
         '  BatchNorm1d after each hidden layer: normalizes activations, helps training.\n'
         '  ReLU activations: standard non-linearity, no vanishing gradient problem.\n'
         '  No activation on output: C/S can be any positive value.'),
        ('75–166', 'def train_model(df)',
         'Training pipeline:\n'
         '  Lines 84–90: Extract features [M, T, r, σ] and targets [norm_price].\n'
         '  Lines 94–96: 80/20 train/test split with random_split().\n'
         '  Line 98: DataLoader with batch_size=1024, shuffle=True.\n'
         '  Lines 101–105: Auto-selects GPU (CUDA or Apple MPS) or CPU.\n'
         '  Line 108: nn.MSELoss — mean squared error on normalized prices.\n'
         '  Line 109: Adam optimizer, lr=0.001.\n'
         '  Lines 116–148: Training loop:\n'
         '    optimizer.zero_grad() — clear gradients.\n'
         '    outputs = model(X_batch) — forward pass.\n'
         '    loss.backward() — compute gradients via autograd.\n'
         '    optimizer.step() — update weights.\n'
         '  Line 163: torch.save(model.state_dict(), "mlp_pricer.pth") — persist weights.'),
        ('170–233', 'def evaluate_model(model, test_loader, device)',
         'Test evaluation + OOD (Out-of-Distribution) tests:\n'
         '  Lines 182–190: Run inference on test set, multiply by S to get actual prices.\n'
         '  Lines 192–193: MAE and RMSE vs BS ground truth.\n'
         '  Lines 202–209: 7 OOD edge cases (Deep OTM, Deep ITM, Very Short/Long expiry,\n'
         '    Extreme vol, Very low/high spot) — tests generalization beyond training range.'),
        ('237–339', 'def test_real_market(model, device, ticker)',
         'Live SPY market test:\n'
         '  Downloads spot price and option chain via yfinance.\n'
         '  Selects ~1-month expiry, filters for liquid options (volume > 0, IV > 0.01).\n'
         '  For each strike: computes BS price and MLP prediction, measures vs mid-price.\n'
         '  Reports MAE for BS and MLP, counts options where MLP beats BS.'),
    ]

    for lines, code, explanation in nn_lines:
        p = doc.add_paragraph()
        r1 = p.add_run(f'Lines {lines}: ')
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        r1.font.size = Pt(11)
        r2 = p.add_run(code)
        r2.font.name = 'Courier New'
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
        for line in explanation.split('\n'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.4)
            p2.add_run(line).font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 10 — lstm.py
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '10. src/pricers/lstm.py — LSTM Volatility Forecaster (Line by Line)', level=1)
    add_para(doc,
        'A hybrid approach: use an LSTM (time-series model) to predict next-day '
        'realized volatility, then plug that prediction into Black-Scholes.',
        italic=True
    )
    doc.add_paragraph()

    lstm_lines = [
        ('13–27', 'def download_data(ticker, start_date, end_date)',
         'Downloads 5 years of daily OHLCV data for SPY from Yahoo Finance.'),
        ('31–125', 'def preprocess_data(df, window_size, vol_window)',
         '  Line 46: Returns = Close.pct_change() — daily log-returns.\n'
         '  Line 49: Realized_Vol = rolling(30).std() × √252 — annualized 30-day vol.\n'
         '  Lines 69–87: Tries to add VIX (fear index) as a 3rd feature. Falls back to\n'
         '    2 features (Returns, Realized_Vol) if VIX fetch fails.\n'
         '  Line 90: Target = next-day realized volatility (shifted by -1).\n'
         '  Line 95: StandardScaler normalizes features to zero mean, unit variance.\n'
         '  Lines 98–102: Creates sliding windows of 60 days (window_size=60).\n'
         '    X[i] = features[i : i+60], y[i] = vol[i+60].\n'
         '  Lines 107–112: Chronological 80/20 split (no shuffling — time series!).'),
        ('129–147', 'class LSTMVolatilityPredictor(nn.Module)',
         '2-layer LSTM with 64 hidden units, dropout=0.2 between layers.\n'
         '  forward(): x is (batch, 60, n_features).\n'
         '    out, _ = lstm(x) — out shape: (batch, 60, 64).\n'
         '    out[:, -1, :] — takes only the last time step\'s hidden state.\n'
         '    linear(out) — maps 64→1 (predicted volatility).'),
        ('151–218', 'def train_model(model, X_train, y_train, epochs, lr, batch_size)',
         'Standard PyTorch training loop with MSE loss and Adam optimizer.\n'
         '  Further 80/20 validation split inside train for overfitting monitoring.\n'
         '  Saves loss_curve.png every epoch.'),
        ('222–261', 'def evaluate_and_compare(model, X_test, y_test, dates_test, hist_vol_test)',
         'Plots LSTM predicted vol vs actual realized vol over the test period.\n'
         'Saves volatility_forecast_comparison.png.'),
        ('265–307', 'def run_bs_experiment(preds, actuals, hist_vol_test)',
         'The key experiment:\n'
         '  For each test day, computes 3 option prices using BS:\n'
         '    1. True price: BS with actual future vol (upper bound).\n'
         '    2. LSTM price: BS with LSTM-predicted vol.\n'
         '    3. Historical price: BS with yesterday\'s 30-day realized vol.\n'
         '  Compares MAE(LSTM-BS, True) vs MAE(Hist-BS, True).\n'
         '  If LSTM forecasts vol better, it prices options more accurately.'),
    ]

    for lines, code, explanation in lstm_lines:
        p = doc.add_paragraph()
        r1 = p.add_run(f'Lines {lines}: ')
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        r1.font.size = Pt(11)
        r2 = p.add_run(code)
        r2.font.name = 'Courier New'
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
        for line in explanation.split('\n'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.4)
            p2.add_run(line).font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 11 — vae.py
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '11. src/pricers/vae.py — Variational Autoencoder (Line by Line)', level=1)
    add_para(doc,
        'Learns the latent structure of the IV surface (5 expiries × 10 strikes = 50D)'
        ' and can generate new surfaces or impute missing data.',
        italic=True
    )
    doc.add_paragraph()

    vae_lines = [
        ('14–28', 'class VAE __init__(input_dim=50, latent_dim=8)',
         'Architecture:\n'
         '  Encoder: 50 → 64 → 32 → (μ, log σ²)  — two parallel linear heads.\n'
         '  Decoder: 8 → 32 → 64 → 50 with Softplus final activation.\n'
         '  Softplus ensures reconstructed IVs are always positive.'),
        ('30–41', 'def encode(x)',
         'Maps the 50-dimensional IV surface to (μ, log σ²) in 8D latent space.\n'
         'Two separate linear layers (fc_mean, fc_logvar) share the 32-D hidden state.'),
        ('43–55', 'def reparameterize(mu, logvar)',
         'The reparameterization trick — allows gradients to flow through sampling:\n'
         '  std = exp(0.5 × logvar)\n'
         '  ε ~ N(0, I)  — random noise\n'
         '  z = μ + ε × std\n'
         '  Instead of sampling z directly (non-differentiable), we sample ε and\n'
         '  compute z deterministically, enabling backpropagation.'),
        ('57–72', 'def decode(z)',
         'Maps 8D latent vector back to 50D IV surface.\n'
         'ReLU activations in hidden layers, Softplus on output (ensures IV > 0).'),
        ('74–86', 'def forward(x)',
         'Full VAE forward pass: encode → reparameterize → decode.\n'
         'Returns (reconstruction, mu, logvar) — all needed for the ELBO loss.'),
        ('89–109', 'def vae_loss_function(recon_x, x, mu, logvar, beta)',
         'ELBO loss = Reconstruction + β × KL Divergence:\n'
         '  MSE = sum((recon_x - x)²)  — how well the VAE reconstructs surfaces.\n'
         '  KLD = -0.5 × sum(1 + logvar - μ² - exp(logvar))  — KL divergence\n'
         '    from the posterior q(z|x) to the prior N(0,I).\n'
         '    Forces the latent space to be compact and well-organized.\n'
         '  β=1 is standard VAE; β > 1 would force more disentanglement.'),
    ]

    for lines, code, explanation in vae_lines:
        p = doc.add_paragraph()
        r1 = p.add_run(f'Lines {lines}: ')
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        r1.font.size = Pt(11)
        r2 = p.add_run(code)
        r2.font.name = 'Courier New'
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
        for line in explanation.split('\n'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.4)
            p2.add_run(line).font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTIONS 12 & 13 — data.py + interpolate.py
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '12. src/utils/data.py — Data Fetching Utilities (Line by Line)', level=1)

    data_lines = [
        ('13–35', 'def get_sp500_tickers()',
         'Returns a hardcoded list of 20 major S&P 500 tickers (AAPL, MSFT, GOOGL…).'),
        ('38–109', 'def fetch_iv_surface(ticker_symbol)',
         '  Fetches up to 10 option expiries for a ticker via yfinance.\n'
         '  For each expiry: downloads call chain, filters IV > 0.01 and bid > 0.\n'
         '  Lines 80–108: Interpolates sparse IV data onto a 5×10 target grid:\n'
         '    5 target DTEs: [30, 60, 90, 120, 150] days.\n'
         '    10 target moneyness levels: linspace(0.8, 1.2, 10).\n'
         '    Uses scipy griddata (linear) for interpolation, nearest for NaN fill.\n'
         '    Returns a (5, 10) NumPy array or None if insufficient data.'),
        ('112–176', 'def main()',
         '  Loops through up to 300 tickers, calls fetch_iv_surface for each.\n'
         '  Lines 139–166: If fewer than 500 real surfaces fetched, generates synthetic ones:\n'
         '    Synthetic surface = base_vol + term_structure·(DTE/30) + skew·(M-1)\n'
         '    + convexity·(M-1)² + noise  — mimics realistic SABR-like smiles.\n'
         '  Saves all 500 surfaces as iv_surfaces.pt.'),
    ]

    for lines, code, explanation in data_lines:
        p = doc.add_paragraph()
        r1 = p.add_run(f'Lines {lines}: ')
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        r1.font.size = Pt(11)
        r2 = p.add_run(code)
        r2.font.name = 'Courier New'
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
        for line in explanation.split('\n'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.4)
            p2.add_run(line).font.size = Pt(11)
        doc.add_paragraph()

    doc.add_paragraph()
    add_heading(doc, '13. src/utils/interpolate.py — IV Surface Interpolation (Line by Line)', level=1)

    interp_lines = [
        ('9–36', 'def impute_vae(model, surface, mask, steps, lr)',
         'VAE-based imputation via latent optimization:\n'
         '  Line 17: z = torch.randn((1, 8), requires_grad=True)  — random starting point in latent space.\n'
         '  Line 18: Adam optimizer on z — we optimize z, not the model weights.\n'
         '  Loop (lines 24–31): For each step:\n'
         '    recon = model.decode(z).view(5,10)  — decode z to 5×10 surface.\n'
         '    loss = MSE only on KNOWN points (mask = False).\n'
         '    Backprop and step — z moves to minimize error on known data.\n'
         '  Line 33: Final decode with best z — fills in the missing cells.'),
        ('39–127', 'def main()',
         '  Loads VAE model and first valid IV surface from iv_surfaces.pt.\n'
         '  Lines 60–63: Creates a missing mask (True = missing) for cells [1:3, 3:7] —\n'
         '    simulates a "chunk" of the surface being unavailable.\n'
         '  Runs both linear interpolation (griddata) and VAE imputation.\n'
         '  Generates 3-panel comparison plot: Original | Linear | VAE.'),
    ]

    for lines, code, explanation in interp_lines:
        p = doc.add_paragraph()
        r1 = p.add_run(f'Lines {lines}: ')
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        r1.font.size = Pt(11)
        r2 = p.add_run(code)
        r2.font.name = 'Courier New'
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
        for line in explanation.split('\n'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.4)
            p2.add_run(line).font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 14 — app.py
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '14. app.py — Streamlit Dashboard (Line by Line)', level=1)
    add_para(doc,
        '1,739-line interactive web application. Brings all pricers together in '
        'a real-time dashboard.',
        italic=True
    )
    doc.add_paragraph()

    app_lines = [
        ('1–9', 'Module docstring',
         'Documents the 4 tabs: Pricing, Greeks, Vol Smile, Model Comparison.'),
        ('11–24', 'Imports',
         'numpy, pandas, streamlit, plotly — UI and data.\n'
         'sys.path.insert — ensures src/ is importable from the root directory.'),
        ('26–44', 'Import pricing functions',
         'Imports black_scholes, Greeks, MC functions, CRR, Heston from their modules.\n'
         'st.set_page_config: wide layout, dark mode compatible.'),
        ('47–202', 'Custom CSS (st.markdown)',
         'Injects CSS variables: dark background (#0d1117), card background (#161b22),\n'
         'accent colors (blue, pink, green, gold). Styles:\n'
         '  metric cards (rounded, hover border glow)\n'
         '  tabs (active = blue background, selected font = black)\n'
         '  price cards (hover lift animation)\n'
         '  buttons (gradient blue-purple)\n'
         '  hero section (gradient top border, multi-column layout).'),
        ('204–244', 'PLOTLY_LAYOUT + COLORS',
         'Shared dark-theme Plotly layout dict applied to every figure.\n'
         'COLORS dict maps model names to their accent colors.'),
        ('246–326', 'Sidebar',
         'All input sliders:\n'
         '  Market Parameters: S, K, T, r, σ.\n'
         '  Option Type: Call/Put selectbox.\n'
         '  Monte Carlo: number of paths (1k/10k/50k/100k).\n'
         '  CRR Binomial: N steps, European/American exercise.\n'
         '  Heston: κ, θ, σᵥ, ρ, v₀, n_paths.'),
        ('329–428', 'Cached compute functions (@st.cache_data)',
         'compute_bs, compute_crr, compute_mc_all, compute_heston,\n'
         'compute_heston_smile_cached, compute_greeks — all decorated with\n'
         '@st.cache_data so they only recompute when inputs change.\n'
         'Pre-computed on every page load using current sidebar values.'),
        ('470–477', 'Tabs definition',
         'st.tabs([Pricing, Greeks, Vol Smile, Model Comparison]) — 4 tabs.'),
        ('482–754', 'Tab 1 — Pricing',
         '  6 price cards (one per model) with color-coded deviation from BS.\n'
         '  Bar chart of all 6 prices with BS reference line.\n'
         '  MC Convergence: 2-panel subplot (price convergence + SE vs N).\n'
         '  Summary table with price, vs-BS diff, SE, speed, type.'),
        ('757–end', 'Tab 2 — Greeks, Tab 3 — Vol Smile, Tab 4 — Model Comparison',
         'Tab 2: BS Greeks bar chart + Delta vs spot + Gamma vs spot plots.\n'
         'Tab 3: Heston IV smile vs flat BS (computed fresh from Heston engine).\n'
         'Tab 4: Loads comparison_table.csv and renders it with conditional coloring.'),
    ]

    for lines, code, explanation in app_lines:
        p = doc.add_paragraph()
        r1 = p.add_run(f'Lines {lines}: ')
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        r1.font.size = Pt(11)
        r2 = p.add_run(code)
        r2.font.name = 'Courier New'
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
        for line in explanation.split('\n'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.4)
            p2.add_run(line).font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 15 — test_greeks.py
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '15. tests/test_greeks.py — Unit Tests (Line by Line)', level=1)
    add_para(doc,
        '91-line unittest suite that mathematically verifies the Greeks '
        'implementation against known financial properties.',
        italic=True
    )
    doc.add_paragraph()

    test_lines = [
        ('1–11', 'Imports + TestCase setup',
         'Imports all 5 Greek functions and implied_vol.\n'
         'setUp: S=K=100, T=1, r=5%, σ=20% — standard ATM parameters.'),
        ('21–41', 'test_delta_bounds_and_atm',
         '  ATM call delta ∈ (0.5, 0.7) — slightly above 0.5 due to continuous drift.\n'
         '  ATM put delta ∈ (-0.5, -0.3) — mirror of call.\n'
         '  call_delta - put_delta = 1.0 exactly (put-call parity for delta).\n'
         '  Deep ITM call delta ≈ 1.0, Deep OTM call delta ≈ 0.0.'),
        ('43–55', 'test_gamma_peaks_atm_and_positive',
         '  Gamma > 0 always (any moneyness).\n'
         '  Gamma(ATM) > Gamma(OTM) and Gamma(ATM) > Gamma(ITM) — peaks at-the-money.'),
        ('57–69', 'test_vega_properties',
         '  Vega > 0 at all moneyness levels.\n'
         '  Vega peaks ATM: Vega(ATM) > Vega(OTM) and Vega(ATM) > Vega(ITM).'),
        ('71–75', 'test_theta_negative_time_decay',
         '  Theta < 0 for both calls and puts.\n'
         '  Options always lose value as time passes (all else equal).'),
        ('77–81', 'test_rho_signs',
         '  Call rho > 0: higher interest rate → call worth more (delayed payment discounted less).\n'
         '  Put rho < 0: higher interest rate → put worth less (strike has lower PV).'),
        ('83–86', 'test_implied_vol_skeleton',
         '  Calls implied_vol(market_price=10, S=K=100, T=1, r=5%) and checks return is not None.\n'
         '  Verifies the function exists, runs, and returns a number.'),
    ]

    for lines, code, explanation in test_lines:
        p = doc.add_paragraph()
        r1 = p.add_run(f'Lines {lines}: ')
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        r1.font.size = Pt(11)
        r2 = p.add_run(code)
        r2.font.name = 'Courier New'
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
        for line in explanation.split('\n'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.4)
            p2.add_run(line).font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 16 — Scripts
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '16. scripts/ — Helper Scripts Explained', level=1)
    script_items = [
        ('train.py', 'Entry-point to train the MLP pricer. Run: python scripts/train.py'),
        ('generate.py', 'Fetches/synthesizes IV surfaces. Run: python scripts/generate.py'),
        ('benchmark_crr.py', 'Plots CRR convergence to BS for N=10→500. Run standalone.'),
        ('pricer_comparison.py', 'Prices one standard option with all 6 models, saves comparison bar chart.'),
        ('volatility_smile.py', 'Generates the Heston IV smile plot. Run standalone.'),
        ('volatility_surface.py', 'Generates 3D IV surface visualization across strikes and expiries.'),
        ('eval_real.py', 'Full evaluation on real SPY options. Reads results_detailed.csv.'),
        ('dm_test.py', 'Diebold-Mariano test: statistically compares two models\' forecast errors.'),
        ('run_definitive_experiment.py',
         '~1,000-line master script. Runs all 7 models on all 200 SPY options, '
         'generates all paper figures and CSV tables. This is the single reproducibility script.'),
        ('generate_pdf.py', 'Assembles and formats the academic paper PDF programmatically.'),
        ('generate_html_pdf.py', 'Alternative: markdown → HTML → PDF conversion pipeline.'),
    ]
    for name, desc in script_items:
        p = doc.add_paragraph()
        r1 = p.add_run(f'📜 {name}: ')
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        p.add_run(desc).font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 17 — END-TO-END DATA FLOW
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '17. How Everything Connects — End-to-End Data Flow', level=1)
    add_para(doc,
        'Here is the complete lifecycle of the project, from raw market data '
        'to the final benchmarks.',
        italic=True
    )
    doc.add_paragraph()

    flow_steps = [
        ('Step 1 — IV Surface Collection', 'data.py',
         'yfinance → 300 tickers → fetch_iv_surface() → 5×10 grid interpolation → '
         'synthetic padding → iv_surfaces.pt (500 surfaces)'),
        ('Step 2 — VAE Training', 'scripts/train.py + vae.py',
         'Load iv_surfaces.pt → VAE(50D→8D→50D) → ELBO loss → vae_weights.pth'),
        ('Step 3 — Synthetic Data Generation', 'neural_net.py',
         '100,000 random (S,K,T,r,σ) → black_scholes() → DataFrame with M = ln(K/S), '
         'norm_price = C/S'),
        ('Step 4 — MLP Training', 'neural_net.py',
         'DataFrame → MLPPricer(4→256→128→64→1) → 100 epochs → mlp_pricer.pth'),
        ('Step 5 — SPY Test Set Collection', 'neural_net.py / run_definitive_experiment.py',
         'yfinance SPY options → 200 contracts → results_detailed.csv '
         '(Strike, Market, IV, T, S, Moneyness)'),
        ('Step 6 — Model Evaluation', 'run_definitive_experiment.py',
         'For each of 200 options, each of 7 models prices it:\n'
         'BS → analytical. CRR → N=200 lattice. MC → 100k paths.\n'
         'LSTM → forecast vol → BS. MLP → mlp_pricer.pth inference.\n'
         'VAE-IV → impute surface → BS interpolation. Heston → calibrate + MC.'),
        ('Step 7 — Metrics & Paper', 'run_definitive_experiment.py + generate_pdf.py',
         'MAE/RMSE/% within 5% per model → comparison_table.csv → LaTeX/PDF paper.'),
        ('Step 8 — Dashboard', 'app.py',
         'streamlit run app.py → user sliders → live pricing + Greeks + Heston smile.'),
    ]

    for step, file, desc in flow_steps:
        p = doc.add_paragraph()
        r1 = p.add_run(f'► {step}  ')
        r1.bold = True
        r1.font.size = Pt(12)
        r1.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        r2 = p.add_run(f'[{file}]')
        r2.italic = True
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        for line in desc.split('\n'):
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.5)
            p2.add_run(line).font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 18 — MATHEMATICAL FOUNDATIONS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '18. Mathematical Foundations', level=1)

    math_sections = [
        ('Black-Scholes Formula',
         'Call: C = S·N(d₁) − K·e^(−rT)·N(d₂)\n'
         'Put:  P = K·e^(−rT)·N(−d₂) − S·N(−d₁)\n\n'
         'd₁ = [ln(S/K) + (r + σ²/2)·T] / (σ·√T)\n'
         'd₂ = d₁ − σ·√T\n\n'
         'N(·) = CDF of standard normal\n'
         'Assumptions: constant volatility, log-normal stock price, '
         'continuous trading, no dividends.'),
        ('Monte Carlo GBM',
         'Under risk-neutral measure Q:\n'
         'S_T = S₀ · exp((r − σ²/2)·T + σ·√T·Z),  Z ~ N(0,1)\n'
         'Price = e^(−rT) · E^Q[max(S_T − K, 0)]\n\n'
         'Standard Error: SE = e^(−rT) · std(payoffs) / √N\n'
         'Converges as O(1/√N) — 10× more paths → ~3× more precision.'),
        ('CRR Binomial Parameters',
         'u = exp(σ·√dt),   d = 1/u,   dt = T/N\n'
         'p = (e^(r·dt) − d) / (u − d)    [risk-neutral probability]\n\n'
         'Backward induction: V(i,j) = disc·[p·V(i+1,j+1) + (1−p)·V(i+1,j)]\n'
         'American: V(i,j) = max(V_cont, Intrinsic)'),
        ('Heston Model',
         'dS = r·S·dt + √v·S·dW₁\n'
         'dv = κ(θ−v)dt + σᵥ·√v·dW₂\n'
         'corr(dW₁, dW₂) = ρ\n\n'
         'κ = mean-reversion speed   θ = long-run variance\n'
         'σᵥ = vol-of-vol   ρ = spot-vol correlation (typically < 0)\n'
         'Negative ρ creates the real-market downside skew.'),
        ('VAE ELBO Loss',
         'L(θ,φ;x) = E_q[log p(x|z)] − KL(q(z|x) || p(z))\n'
         '          = −MSE(x, decode(z)) − 0.5·Σ(1 + log σ² − μ² − σ²)\n'
         'p(z) = N(0,I)  [standard normal prior]\n'
         'q(z|x) = N(μ(x), diag(σ²(x)))  [encoder posterior]'),
    ]

    for title, content in math_sections:
        add_para(doc, title, bold=True, size=12, color='1a56db')
        add_code_block(doc, content)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 19 — EMPIRICAL RESULTS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '19. Empirical Results & Benchmarks', level=1)
    add_para(doc,
        'Evaluated on 200 real SPY options. Metrics: MAE (mean absolute error '
        'in dollars), RMSE, % priced within 5% of market, ATM MAE, OTM MAE, '
        'and speed (milliseconds per option).',
    )
    doc.add_paragraph()

    # Table
    table = doc.add_table(rows=9, cols=8)
    table.style = 'Light Grid Accent 1'

    headers = ['Method', 'MAE', 'RMSE', '% ≤5%', 'ATM MAE', 'OTM MAE', 'Speed (ms)', 'Best Use']
    rows_data = [
        ['Black-Scholes',    '0.9058', '1.4442', '49.0%', '0.3790', '0.0193', '0.0096',  'Fast European baseline'],
        ['CRR Binomial',     '0.9053', '1.4431', '47.5%', '0.3793', '0.0198', '13.5539', 'American / early exercise'],
        ['Monte Carlo 100k', '0.9038', '1.4459', '50.5%', '0.3808', '0.0199', '1.2256',  'Exotic / path-dependent'],
        ['LSTM-BS Hybrid',   '1.0477', '1.4902', '41.0%', '1.2547', '0.0413', '0.0094',  'Vol forecasting'],
        ['MLP Pricer',       '6.6577', '7.3106', '12.5%', '10.9255','5.1214', '0.4861',  'Ultra-fast batch pricing'],
        ['VAE-IV → BS',      '2.4755', '3.1479', '25.5%', '4.2988', '0.9144', '0.0347',  'IV surface interpolation'],
        ['Heston MC',        '1.0402', '1.5781', '38.5%', '0.4550', '0.1459', '18.2100', 'Stochastic vol / smile'],
    ]

    # Header row
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        set_cell_background(cell, '1a56db')
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)

    # Data rows
    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if r_idx % 2 == 0:
                set_cell_background(cell, 'f0f4ff')

    doc.add_paragraph()
    add_para(doc,
        'Key insight: Black-Scholes, CRR, and Monte Carlo all perform similarly '
        '(MAE ~0.90) on this test set — because SPY options are liquid and well-priced. '
        'Heston adds smile-consistency at the cost of 18ms/option. MLP is fastest but '
        'least accurate on real market data (trained on synthetic BS prices, not market prices).',
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 20 — ADVANTAGES
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '20. Advantages of the Project', level=1)

    advantages = [
        ('Comprehensive methodology coverage',
         'Seven distinct pricing approaches in one codebase — from 1973 Black-Scholes '
         'to 2024 deep learning. Rare to find all in one project.'),
        ('Research-grade implementation',
         'Hand-coded norm_cdf (Abramowitz & Stegun), proper variance reduction techniques '
         '(antithetic variates, control variates), full-truncation Euler-Maruyama for Heston — '
         'not just wrappers around libraries.'),
        ('Real market validation',
         '200 live SPY options from Yahoo Finance provide a credible benchmark. '
         'Results are publishable (submitted to SSRN).'),
        ('Complete academic paper',
         'The project produces a full LaTeX/PDF paper with methodology, results tables, '
         'and figures — going beyond a typical coding project.'),
        ('Interactive Streamlit dashboard',
         'Any parameter can be changed in real time via sliders. All 6 models '
         'recompute and display price + Greeks + vol smile simultaneously.'),
        ('Auto-differentiation Greeks',
         'mlp_greeks.py demonstrates computing neural network derivatives via autograd — '
         'a technique used in production derivatives desks.'),
        ('Exotic option support',
         'mc_barrier_option() prices path-dependent barrier options — impossible '
         'with closed-form BS, handled naturally by path simulation.'),
        ('Variance reduction research',
         'The Monte Carlo module demonstrates all three main variance reduction techniques '
         'with convergence plots and statistical analysis.'),
        ('Scale-invariant neural design',
         'The MLP learns C/S = f(M, T, r, σ) — scale invariant in spot price, '
         'meaning it generalizes to any stock level.'),
        ('Reproducible research',
         'Fixed random seeds, pinned requirements.txt, and the master '
         'run_definitive_experiment.py script make all results exactly reproducible.'),
    ]

    for i, (title, desc) in enumerate(advantages, 1):
        p = doc.add_paragraph()
        r1 = p.add_run(f'✅ {i}. {title}: ')
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(0x05, 0x73, 0x43)
        p.add_run(desc).font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 21 — LIMITATIONS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '21. Limitations of the Project', level=1)

    limitations = [
        ('MLP trained on synthetic BS data',
         'The MLP learns to replicate Black-Scholes, not real market prices. '
         'When real prices deviate from BS (smile, skew), the MLP inherits BS\'s errors. '
         'MAE = 6.66 on real data is poor.'),
        ('Heston calibration is slow and noisy',
         'Calibrating 5 Heston parameters per option using Nelder-Mead with '
         'only 2,000 MC paths produces noisy estimates. Each calibration takes ~18ms.'),
        ('No Greeks for MC or Heston',
         'Monte Carlo and Heston don\'t produce Delta/Gamma analytically. '
         'Finite-difference Greeks for these models are slow and noisy.'),
        ('Heston Euler-Maruyama discretization bias',
         'Full-truncation Euler-Maruyama introduces bias. '
         'More accurate schemes (Broadie-Kaya exact, QE scheme) are not implemented.'),
        ('LSTM uses rolling realized vol, not implied vol',
         'The LSTM forecasts realized (historical) volatility. '
         'Real traders use implied volatility as the primary input. '
         'This is a conceptual mismatch.'),
        ('No dividend handling',
         'All models assume zero dividends. SPY pays ~1.3% annual dividend, '
         'causing systematic underpricing of long-dated calls.'),
        ('No transaction costs or bid-ask spreads',
         'Models compare to the mid-price (bid+ask)/2. In practice, the bid-ask '
         'spread is 1–5% of price for liquid options, dwarfing model errors.'),
        ('Single test date',
         'The 200 SPY options are from one day. Results may not generalize '
         'to different market regimes (high-vol periods, market crashes).'),
        ('Heston duplicate function definition',
         'In monte_carlo.py, mc_barrier_option() is defined twice (lines 224 and 291). '
         'The first definition has a pass statement. Python silently uses the second. '
         'This is a code quality issue.'),
        ('VAE trained on partially synthetic data',
         'Only real IV surfaces where all 10 strikes × 5 expiries have valid data are kept. '
         'The remainder is filled with synthetic SABR-like surfaces. '
         'This may bias the VAE toward idealized surfaces.'),
    ]

    for i, (title, desc) in enumerate(limitations, 1):
        p = doc.add_paragraph()
        r1 = p.add_run(f'⚠️ {i}. {title}: ')
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor(0xb4, 0x53, 0x09)
        p.add_run(desc).font.size = Pt(11)
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 22 — IMPROVEMENT ROADMAP
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '22. How to Improve the Project — Roadmap', level=1)
    add_para(doc,
        'Prioritized improvement ideas from quick wins to major research directions.',
        italic=True
    )
    doc.add_paragraph()

    improvements = [
        ('QUICK WINS', [
            ('Fix duplicate mc_barrier_option()',
             'Remove the first (broken) definition of mc_barrier_option() in monte_carlo.py lines 224–289.'),
            ('Add dividend yield parameter',
             'All pricers should accept a dividend yield q. '
             'BS call becomes: C = S·e^(−qT)·N(d₁) − K·e^(−rT)·N(d₂). '
             'Only a 2-line change per pricer.'),
            ('Add more unit tests',
             'tests/ only has one file. Add tests for:\n'
             '  • MC convergence (BS error < 0.01 at N=100k)\n'
             '  • CRR convergence (error < 0.001 at N=500)\n'
             '  • Heston sanity check (σᵥ→0, ρ→0 → Heston ≈ BS)\n'
             '  • Put-call parity for all pricers'),
            ('Add American put to Heston',
             'Current Heston only prices European options. Add Longstaff-Schwartz '
             'LSM regression for American exercise within the Heston MC paths.'),
        ]),
        ('MEDIUM IMPROVEMENTS', [
            ('Train MLP on real market prices',
             'Retrain the MLP using real bid-ask midpoints from options chains '
             '(not BS-generated prices). This will teach the model the actual vol smile '
             'and dramatically reduce MAE from 6.66 to potentially < 1.0.'),
            ('Use implied vol as LSTM input',
             'Replace the LSTM\'s realized vol target with implied vol (VIX or at-the-money IV). '
             'Implied vol is forward-looking and more relevant for pricing.'),
            ('Implement Heston QE scheme',
             'Replace Euler-Maruyama with the Quadratic-Exponential (QE) discretization '
             '(Andersen 2007) — eliminates discretization bias with similar computational cost.'),
            ('Add Quasi-Monte Carlo (Sobol sequences)',
             'Replace standard normal draws in MC with Sobol or Halton low-discrepancy sequences. '
             'Reduces MC error from O(1/√N) to O(log(N)^d / N) — huge improvement.'),
            ('Greeks for all models via finite difference',
             'Add a generic numerical_greeks() function using 2-point finite difference '
             '(bump by ε, compute price, divide by ε). Applicable to any pricer.'),
            ('Add SABR model',
             'SABR (Hagan 2002) is the industry standard for smile interpolation. '
             'Has a semi-analytical formula. Much faster than Heston MC.'),
        ]),
        ('MAJOR RESEARCH DIRECTIONS', [
            ('Train MLP end-to-end on real options data',
             'Collect 2+ years of daily SPY options chains (~500k+ data points). '
             'Train MLP to directly predict market price from (S, K, T, r, IV). '
             'This could achieve MAE < 0.10.'),
            ('Transformer-based pricer',
             'Replace the MLP with a Transformer that takes the full option chain '
             '(all strikes × expiries) as input and outputs prices for all options simultaneously. '
             'This captures cross-strike and cross-expiry relationships.'),
            ('Online learning / live dashboard',
             'Connect the dashboard to a live options data feed (Polygon.io, Tradier). '
             'Retrain or fine-tune the ML models online as new quotes arrive.'),
            ('Local volatility model (Dupire)',
             'Dupire (1994) derives the exact local volatility surface from market option prices. '
             'Add dupire_local_vol() to src/pricers/ as another benchmark.'),
            ('Risk management integration',
             'Add portfolio-level risk metrics:\n'
             '  • Delta/Vega/Gamma PnL attribution\n'
             '  • Value at Risk (VaR) via MC on a portfolio\n'
             '  • Stress testing (what-if scenarios)'),
            ('Publish on PyPI as a library',
             'Restructure the project as an installable Python package (pyproject.toml). '
             'Add type stubs, docstrings, and CI/CD. Publish to PyPI as options-pricer. '
             'This makes the project usable by other quants.'),
        ]),
    ]

    for category, items in improvements:
        add_para(doc, category, bold=True, size=13, color='1a56db')
        for title, desc in items:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            r1 = p.add_run(f'▶ {title}: ')
            r1.bold = True
            r1.font.size = Pt(11)
            for line in desc.split('\n'):
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.6)
                p2.add_run(line).font.size = Pt(11)
            doc.add_paragraph()

    # ── Final page ────────────────────────────────────────────────────────────
    doc.add_page_break()
    final_p = doc.add_paragraph()
    final_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = final_p.add_run('— End of Documentation —')
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.add_run(
        'Options Pricing Engine · Chaithanya · 2026\n'
        'github.com/Chaithanya5gif/Options-pricing-engine'
    ).font.size = Pt(11)

    # ── Save ──────────────────────────────────────────────────────────────────
    output_path = 'Options_Pricing_Engine_Documentation.docx'
    doc.save(output_path)
    print(f'✅ Documentation saved to: {output_path}')
    return output_path


if __name__ == '__main__':
    build_doc()
